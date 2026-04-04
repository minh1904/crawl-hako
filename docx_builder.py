"""
docx_builder.py — Tạo file DOCX từ danh sách chapter
"""
import io
import logging
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

logger = logging.getLogger(__name__)


def build_docx(
    output_path: Path,
    novel_info: dict,
    volume_title: str,
    chapters_data: list[dict],
    cover_bytes: bytes | None,
    image_cache: dict[str, bytes],
) -> None:
    """
    Tạo file DOCX.

    chapters_data: list of {title, elements: [{type, content|url}]}
    image_cache: {url: bytes}
    """
    doc = Document()

    # Cài font mặc định
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # Trang bìa
    _add_cover_page(doc, novel_info, volume_title, cover_bytes)

    for chap_idx, chap in enumerate(chapters_data):
        chap_title = chap.get("title", f"Chương {chap_idx + 1}")
        elements = chap.get("elements", [])

        # Page break trước mỗi chương (trừ chương đầu tiên vì đã có trang bìa)
        if chap_idx > 0:
            doc.add_page_break()

        # Tiêu đề chương
        h = doc.add_heading(chap_title, level=2)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for elem in elements:
            if elem["type"] == "text":
                text = elem["content"]
                if text:
                    para = doc.add_paragraph(text)
                    para.paragraph_format.first_line_indent = Inches(0.3)
                else:
                    doc.add_paragraph("")
            elif elem["type"] == "image":
                url = elem["url"]
                img_bytes = image_cache.get(url)
                if img_bytes:
                    _add_image(doc, img_bytes, url)
                else:
                    doc.add_paragraph(f"[Ảnh: {url}]").italic = True

    doc.save(str(output_path))
    logger.info(f"Đã xuất DOCX: {output_path}")


def _add_cover_page(doc: Document, novel_info: dict, volume_title: str, cover_bytes: bytes | None) -> None:
    """Thêm trang bìa vào đầu document."""
    title = novel_info.get("title", "Unknown")
    author = novel_info.get("author", "")

    if cover_bytes:
        _add_image(doc, cover_bytes, "cover", max_width_inches=5.0)

    # Tên truyện
    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Tên tập
    if volume_title:
        vol_para = doc.add_paragraph(volume_title)
        vol_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        vol_run = vol_para.runs[0] if vol_para.runs else vol_para.add_run(volume_title)
        vol_run.bold = True
        vol_run.font.size = Pt(14)

    # Tác giả
    if author:
        auth_para = doc.add_paragraph(f"Tác giả: {author}")
        auth_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()


def _add_image(doc: Document, img_bytes: bytes, label: str, max_width_inches: float = 4.5) -> None:
    """Thêm ảnh vào document, tự co giãn theo max_width."""
    try:
        pil_img = Image.open(io.BytesIO(img_bytes))
        # Convert WEBP → JPEG
        if pil_img.format == "WEBP" or pil_img.mode in ("RGBA", "P"):
            buf = io.BytesIO()
            pil_img.convert("RGB").save(buf, format="JPEG", quality=90)
            img_bytes = buf.getvalue()

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(io.BytesIO(img_bytes), width=Inches(max_width_inches))
    except Exception as e:
        logger.warning(f"Không thêm được ảnh ({label}): {e}")
        doc.add_paragraph(f"[Không hiển thị được ảnh: {label}]")
